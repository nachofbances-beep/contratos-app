"""
Genera las plantillas .docx de los contratos.
Solo necesitas ejecutar este script UNA VEZ, al instalar la app o cuando
añadas nuevos modelos:

  python crear_plantilla.py

Después podrás abrir los archivos en /modelos/ con Word y modificar el
texto a tu gusto, siempre conservando los marcadores {{ campo }}.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def add_titulo(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(14)
    return p


def add_seccion(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    return p


def add_clausula(doc, titulo, texto):
    p = doc.add_paragraph()
    run = p.add_run(titulo)
    run.bold = True
    p.add_run(" " + texto)
    return p


def crear_compraventa_inmueble(ruta_destino):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_titulo(doc, "CONTRATO DE COMPRAVENTA DE BIEN INMUEBLE")
    doc.add_paragraph()

    doc.add_paragraph("En {{ lugar_firma }}, a {{ fecha_firma }}.")
    doc.add_paragraph()

    add_seccion(doc, "REUNIDOS")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("DE UNA PARTE, como ")
    p.add_run("PARTE VENDEDORA").bold = True
    p.add_run(":")
    doc.add_paragraph(
        "D./Dña. {{ vendedor_nombre }}, mayor de edad, con DNI/NIE nº "
        "{{ vendedor_dni }}, y domicilio en {{ vendedor_domicilio }}."
    )
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("DE OTRA PARTE, como ")
    p.add_run("PARTE COMPRADORA").bold = True
    p.add_run(":")
    doc.add_paragraph(
        "D./Dña. {{ comprador_nombre }}, mayor de edad, con DNI/NIE nº "
        "{{ comprador_dni }}, y domicilio en {{ comprador_domicilio }}."
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "Ambas partes se reconocen mutuamente capacidad legal suficiente para "
        "el otorgamiento del presente contrato y, a tal efecto,"
    )
    doc.add_paragraph()

    add_seccion(doc, "EXPONEN")
    doc.add_paragraph()

    doc.add_paragraph(
        "I.- Que la PARTE VENDEDORA es propietaria en pleno dominio del "
        "siguiente bien inmueble: {{ inmueble_descripcion }}, sito en "
        "{{ inmueble_direccion }}. Inscrito en el Registro de la Propiedad "
        "de {{ registro_propiedad }}, al Tomo {{ registro_tomo }}, Libro "
        "{{ registro_libro }}, Folio {{ registro_folio }}, Finca nº "
        "{{ registro_finca }}. Referencia catastral: "
        "{{ referencia_catastral }}."
    )
    doc.add_paragraph(
        "II.- Que el inmueble se encuentra libre de cargas, gravámenes y "
        "arrendatarios, y al corriente de pago de cuotas de comunidad, IBI "
        "y demás impuestos."
    )
    doc.add_paragraph(
        "III.- Que la PARTE COMPRADORA está interesada en adquirir dicho "
        "inmueble, y la PARTE VENDEDORA en transmitirlo, conforme a las "
        "siguientes"
    )
    doc.add_paragraph()

    add_seccion(doc, "CLÁUSULAS")
    doc.add_paragraph()

    add_clausula(
        doc,
        "PRIMERA.- OBJETO.",
        "La PARTE VENDEDORA vende y transmite a la PARTE COMPRADORA, que "
        "adquiere, el inmueble descrito en el Expositivo I del presente "
        "contrato, con todos sus elementos integrantes, accesorios y anejos.",
    )
    add_clausula(
        doc,
        "SEGUNDA.- PRECIO Y FORMA DE PAGO.",
        "El precio total de la presente compraventa se fija en "
        "{{ precio_compra_numero }} euros ({{ precio_compra_letras }}), "
        "cantidad que la PARTE COMPRADORA satisface a la PARTE VENDEDORA "
        "mediante {{ forma_pago }}, sirviendo el presente documento como "
        "la mas eficaz carta de pago.",
    )
    add_clausula(
        doc,
        "TERCERA.- GASTOS E IMPUESTOS.",
        "Los gastos e impuestos derivados de la presente compraventa se "
        "satisfaran conforme a lo dispuesto en la legislacion vigente, "
        "salvo pacto expreso en contrario entre las partes.",
    )
    add_clausula(
        doc,
        "CUARTA.- ENTREGA Y POSESIÓN.",
        "La entrega de la posesión del inmueble se efectuará el día "
        "{{ fecha_entrega }}, mediante la firma de la correspondiente "
        "escritura pública de compraventa ante Notario.",
    )
    add_clausula(
        doc,
        "QUINTA.- SANEAMIENTO.",
        "La PARTE VENDEDORA responde frente a la PARTE COMPRADORA del "
        "saneamiento por evicción y por vicios o defectos ocultos en los "
        "términos previstos en los artículos 1474 y siguientes del Código "
        "Civil.",
    )
    add_clausula(
        doc,
        "SEXTA.- JURISDICCIÓN.",
        "Para cualquier cuestión litigiosa derivada del presente contrato, "
        "las partes se someten expresamente a los Juzgados y Tribunales de "
        "{{ jurisdiccion }}, con renuncia a su propio fuero si lo tuvieren.",
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "Y en prueba de conformidad con cuanto antecede, firman ambas "
        "partes el presente contrato por duplicado y a un solo efecto, en "
        "el lugar y fecha indicados al inicio."
    )
    doc.add_paragraph()
    doc.add_paragraph()

    doc.add_paragraph("LA PARTE VENDEDORA\t\t\t\tLA PARTE COMPRADORA")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_____________________\t\t\t_____________________")
    doc.add_paragraph("{{ vendedor_nombre }}\t\t\t{{ comprador_nombre }}")

    os.makedirs(os.path.dirname(ruta_destino), exist_ok=True)
    doc.save(ruta_destino)
    print(f"Plantilla creada en: {ruta_destino}")


def crear_nda_confidencialidad(ruta_destino):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_titulo(doc, "ACUERDO DE CONFIDENCIALIDAD")
    doc.add_paragraph()

    doc.add_paragraph("En {{ lugar_firma }}, a {{ fecha_firma }}.")
    doc.add_paragraph()

    add_seccion(doc, "REUNIDOS")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("DE UNA PARTE: ").bold = True
    p.add_run(
        "{{ parte_a_nombre }}, con DNI/NIE/CIF nº {{ parte_a_dni }} y "
        "domicilio en {{ parte_a_domicilio }}{{ parte_a_representante }} "
        "(en adelante, la "
    )
    p.add_run('"PARTE A"').bold = True
    p.add_run(").")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("DE OTRA PARTE: ").bold = True
    p.add_run(
        "{{ parte_b_nombre }}, con DNI/NIE/CIF nº {{ parte_b_dni }} y "
        "domicilio en {{ parte_b_domicilio }}{{ parte_b_representante }} "
        "(en adelante, la "
    )
    p.add_run('"PARTE B"').bold = True
    p.add_run(").")
    doc.add_paragraph()

    doc.add_paragraph(
        "Las partes se reconocen mutuamente capacidad legal suficiente "
        "para el otorgamiento del presente Acuerdo y, a tal efecto,"
    )
    doc.add_paragraph()

    add_seccion(doc, "EXPONEN")
    doc.add_paragraph()

    doc.add_paragraph(
        "I.- Que las partes están manteniendo conversaciones y/o "
        "relaciones de negocio en relación con {{ objeto_relacion }} "
        '(en adelante, la "Finalidad").'
    )
    doc.add_paragraph(
        "II.- Que en el marco de dicha Finalidad las partes pueden "
        "intercambiar información de carácter confidencial cuya "
        "divulgación a terceros podría causar perjuicios."
    )
    doc.add_paragraph(
        "III.- Que las partes acuerdan suscribir el presente Acuerdo de "
        "Confidencialidad conforme a las siguientes"
    )
    doc.add_paragraph()

    add_seccion(doc, "CLÁUSULAS")
    doc.add_paragraph()

    add_clausula(
        doc,
        "PRIMERA.- OBJETO.",
        "El presente Acuerdo regula las condiciones bajo las cuales las "
        "partes intercambiarán información confidencial en el marco de "
        "la Finalidad descrita en el Expositivo I.",
    )
    add_clausula(
        doc,
        "SEGUNDA.- INFORMACIÓN CONFIDENCIAL.",
        'Se considerará "Información Confidencial" toda información de '
        "cualquier naturaleza —técnica, comercial, financiera, "
        "organizativa, de personal, datos de clientes o proveedores— que "
        'una parte (la "Parte Divulgadora") revele a la otra (la '
        '"Parte Receptora"), tanto verbalmente como por escrito, en '
        "formato físico o electrónico, con independencia de que esté "
        "marcada o no como confidencial. Quedan excluidas: (i) la "
        "información que ya fuera de dominio público en el momento de su "
        "revelación; (ii) la que pase a serlo posteriormente sin culpa "
        "de la Parte Receptora; (iii) la que la Parte Receptora ya "
        "conociera con anterioridad de forma lícita; y (iv) la que deba "
        "revelarse en cumplimiento de una orden judicial o "
        "administrativa.",
    )
    add_clausula(
        doc,
        "TERCERA.- OBLIGACIONES.",
        "La Parte Receptora se obliga a: (i) mantener la Información "
        "Confidencial en estricto secreto; (ii) utilizarla exclusivamente "
        "para la Finalidad pactada; (iii) no revelarla a terceros sin "
        "consentimiento previo y por escrito de la Parte Divulgadora; "
        "(iv) limitar el acceso a aquellos empleados o colaboradores que "
        "necesiten conocerla, exigiéndoles el mismo deber de "
        "confidencialidad; y (v) emplear el mismo grado de diligencia "
        "que aplica a su propia información confidencial, sin que sea "
        "inferior a una diligencia razonable.",
    )
    add_clausula(
        doc,
        "CUARTA.- DURACIÓN.",
        "El presente Acuerdo entrará en vigor en la fecha de su firma y "
        "permanecerá vigente durante {{ duracion_acuerdo }}. La "
        "obligación de confidencialidad subsistirá durante "
        "{{ plazo_confidencialidad }} años desde la finalización del "
        "Acuerdo, con independencia del motivo de su terminación.",
    )
    add_clausula(
        doc,
        "QUINTA.- PROPIEDAD.",
        "La Información Confidencial es y permanecerá de la exclusiva "
        "propiedad de la Parte Divulgadora. La firma del presente "
        "Acuerdo no implica la concesión de licencia, derecho o "
        "autorización alguna sobre dicha información, salvo lo "
        "expresamente previsto para la Finalidad.",
    )
    add_clausula(
        doc,
        "SEXTA.- DEVOLUCIÓN.",
        "A la finalización del Acuerdo, o en cualquier momento a "
        "requerimiento de la Parte Divulgadora, la Parte Receptora "
        "deberá devolver o destruir toda la Información Confidencial en "
        "su poder, así como cualquier copia o derivado, certificándolo "
        "por escrito si así se le solicita.",
    )
    add_clausula(
        doc,
        "SÉPTIMA.- INCUMPLIMIENTO.",
        "El incumplimiento de las obligaciones aquí asumidas dará "
        "derecho a la parte perjudicada a exigir una indemnización por "
        "los daños y perjuicios causados, fijándose una penalización "
        "mínima de {{ penalizacion }} euros, sin perjuicio de las "
        "acciones legales adicionales que correspondan.",
    )
    add_clausula(
        doc,
        "OCTAVA.- LEY APLICABLE Y JURISDICCIÓN.",
        "El presente Acuerdo se rige por la legislación española. Para "
        "cualquier controversia derivada del mismo, las partes se "
        "someten expresamente a los Juzgados y Tribunales de "
        "{{ jurisdiccion }}, con renuncia a su propio fuero si lo "
        "tuvieren.",
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "Y en prueba de conformidad con cuanto antecede, firman ambas "
        "partes el presente Acuerdo por duplicado y a un solo efecto, "
        "en el lugar y fecha indicados al inicio."
    )
    doc.add_paragraph()
    doc.add_paragraph()

    doc.add_paragraph("LA PARTE A\t\t\t\t\tLA PARTE B")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_____________________\t\t\t_____________________")
    doc.add_paragraph("{{ parte_a_nombre }}\t\t\t{{ parte_b_nombre }}")

    os.makedirs(os.path.dirname(ruta_destino), exist_ok=True)
    doc.save(ruta_destino)
    print(f"Plantilla creada en: {ruta_destino}")


if __name__ == "__main__":
    base = os.path.dirname(__file__)
    crear_compraventa_inmueble(os.path.join(base, "modelos", "compraventa_inmueble.docx"))
    crear_nda_confidencialidad(os.path.join(base, "modelos", "nda_confidencialidad.docx"))
